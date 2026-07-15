"""
Word Generator module for the AI-Powered Smart Recipe Recommendation System.
Generates structured Microsoft Word (.docx) files of recipes with styling.
"""

from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import EXPORTS_DIR
from utils import logger, format_duration


def generate_recipe_docx(recipe: Dict[str, Any]) -> str:
    """
    Generates a Word Document (.docx) file for the given recipe.

    Args:
        recipe (Dict[str, Any]): Recipe data matching our schema.

    Returns:
        str: Absolute path to the generated DOCX.
    """
    try:
        doc = Document()

        # Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Style colors
        color_primary = RGBColor(255, 87, 34)   # Orange
        color_secondary = RGBColor(76, 175, 80) # Green
        color_dark = RGBColor(33, 33, 33)       # Dark Gray
        color_accent = RGBColor(156, 39, 176)   # Purple
        color_blue = RGBColor(33, 150, 243)     # Blue

        # 1. Yemek Adı (Title)
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_title = title_paragraph.add_run(recipe["yemek"])
        run_title.font.name = "Arial"
        run_title.font.size = Pt(24)
        run_title.font.bold = True
        run_title.font.color.rgb = color_primary

        # 2. Açıklama (Description)
        desc_paragraph = doc.add_paragraph()
        run_desc = desc_paragraph.add_run(recipe["aciklama"])
        run_desc.font.name = "Arial"
        run_desc.font.size = Pt(11)
        run_desc.font.italic = True
        run_desc.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph() # Spacer

        # 3. Bilgiler Tablosu (Info Table)
        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        
        # Headers
        headers = ["Hazırlama Süresi", "Zorluk Seviyesi", "Yaklaşık Kalori"]
        for idx, h in enumerate(headers):
            cell = table.cell(0, idx)
            cell.text = h
            # Format header
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"

        # Values
        values = [
            format_duration(recipe["sure_dakika"]),
            recipe["zorluk"],
            f"{recipe['kalori']} kcal"
        ]
        for idx, v in enumerate(values):
            cell = table.cell(1, idx)
            cell.text = v
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(11)
            run.font.name = "Arial"

        doc.add_paragraph() # Spacer

        # 4. Malzemeler (Ingredients)
        heading_ingredients = doc.add_paragraph()
        run_ing_h = heading_ingredients.add_run("📋 Kullanılacak Malzemeler")
        run_ing_h.font.bold = True
        run_ing_h.font.size = Pt(14)
        run_ing_h.font.name = "Arial"
        run_ing_h.font.color.rgb = color_secondary

        for malzeme in recipe["malzemeler"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(malzeme)
            run.font.name = "Arial"
            run.font.size = Pt(10.5)

        # Eksik Malzemeler
        if recipe.get("eksik_malzemeler"):
            doc.add_paragraph() # Spacer
            heading_missing = doc.add_paragraph()
            run_mis_h = heading_missing.add_run("⚠️ Eksik Malzemeler (Gerekli Alışveriş Listesi)")
            run_mis_h.font.bold = True
            run_mis_h.font.size = Pt(12)
            run_mis_h.font.name = "Arial"
            run_mis_h.font.color.rgb = RGBColor(231, 76, 60)

            for eksik in recipe["eksik_malzemeler"]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(eksik)
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(120, 30, 30)

        doc.add_paragraph() # Spacer

        # 5. Yapılış Adımları (Steps)
        heading_steps = doc.add_paragraph()
        run_steps_h = heading_steps.add_run("👨‍🍳 Adım Adım Yapılışı")
        run_steps_h.font.bold = True
        run_steps_h.font.size = Pt(14)
        run_steps_h.font.name = "Arial"
        run_steps_h.font.color.rgb = color_primary

        for idx, adim in enumerate(recipe["yapilis"], 1):
            p = doc.add_paragraph()
            # Numbered layout
            run_num = p.add_run(f"{idx}. ")
            run_num.font.bold = True
            run_num.font.name = "Arial"
            run_num.font.size = Pt(10.5)
            run_num.font.color.rgb = color_primary

            run_text = p.add_run(adim)
            run_text.font.name = "Arial"
            run_text.font.size = Pt(10.5)

        doc.add_paragraph() # Spacer

        # 6. Sunum Önerisi (Presentation)
        heading_pres = doc.add_paragraph()
        run_pres_h = heading_pres.add_run("🍽️ Sunum Önerisi")
        run_pres_h.font.bold = True
        run_pres_h.font.size = Pt(14)
        run_pres_h.font.name = "Arial"
        run_pres_h.font.color.rgb = color_accent

        p_pres = doc.add_paragraph()
        run_pres = p_pres.add_run(recipe["sunum_onerisi"])
        run_pres.font.name = "Arial"
        run_pres.font.size = Pt(10.5)

        doc.add_paragraph() # Spacer

        # 7. Alternatif Tarifler (Alternatives)
        heading_alt = doc.add_paragraph()
        run_alt_h = heading_alt.add_run("🥗 Alternatif Tarif Önerileri")
        run_alt_h.font.bold = True
        run_alt_h.font.size = Pt(14)
        run_alt_h.font.name = "Arial"
        run_alt_h.font.color.rgb = color_blue

        for alt in recipe["alternatif_tarifler"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(alt)
            run.font.name = "Arial"
            run.font.size = Pt(10.5)

        # Document saving
        safe_name = "".join([c if c.isalnum() else "_" for c in recipe["yemek"]])
        docx_filename = f"tarif_{safe_name}.docx"
        file_path = EXPORTS_DIR / docx_filename

        doc.save(str(file_path))
        logger.info(f"DOCX generated successfully at: {file_path}")
        return str(file_path)

    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise e
